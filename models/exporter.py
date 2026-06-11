from docx import Document
from database.database import get_connection

def exportar_faltantes_docx(ruta_salida="faltantes.docx"):
    doc = Document()
    doc.add_heading("AlbumMind — Figuritas Faltantes", level=1)

    with get_connection() as conn:
        faltantes = conn.execute("""
            SELECT f.numFigura, p.nombre AS pais, j.nombre, j.apellido
            FROM FIGURITA f
            JOIN PAIS p ON f.codPais = p.codPais
            LEFT JOIN FJUGADOR j ON (f.numFigura = j.numFigura AND f.codPais = j.codPais)
            LEFT JOIN COLECCION c ON (f.numFigura = c.numFigura AND f.codPais = c.codPais)
            WHERE c.numFigura IS NULL OR c.cantidad = 0
            ORDER BY p.nombre, f.numFigura
        """).fetchall()

    for row in faltantes:
        doc.add_paragraph(
            f"#{row['numFigura']} — {row['pais']}: {row['nombre']} {row['apellido'] or ''}".strip()
        )
    doc.save(ruta_salida)

def exportar_repetidas_docx(ruta_salida="repetidas.docx"):
    doc = Document()
    doc.add_heading("AlbumMind — Figuritas Repetidas (para intercambio)", level=1)

    with get_connection() as conn:
        repetidas = conn.execute("""
            SELECT c.numFigura, p.nombre AS pais, j.nombre, j.apellido,
                   (c.cantidad - 1) AS disponibles
            FROM COLECCION c
            JOIN FIGURITA f ON (c.numFigura = f.numFigura AND c.codPais = f.codPais)
            JOIN PAIS p ON f.codPais = p.codPais
            LEFT JOIN FJUGADOR j ON (c.numFigura = j.numFigura AND c.codPais = j.codPais)
            WHERE c.cantidad > 1
            ORDER BY p.nombre, c.numFigura
        """).fetchall()

    for row in repetidas:
        doc.add_paragraph(
            f"#{row['numFigura']} — {row['pais']}: {row['nombre']} {row['apellido'] or ''} "
            f"(x{row['disponibles']} disponibles)"
        )
    doc.save(ruta_salida)